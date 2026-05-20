import io
import json
import time
import threading
from pathlib import Path

from django.conf import settings
from django.http import StreamingHttpResponse, FileResponse, HttpResponse
from rest_framework import status
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Tender
from .serializers import TenderListSerializer, TenderDetailSerializer, TenderUploadSerializer
from ai_pipeline.pipeline import process_tender
from ai_pipeline.monopoly import check_monopoly, prepare_tender_data
from ai_pipeline.goszakup import search_tenders as goszakup_search, fetch_tender_pdf


@api_view(['GET'])
def tender_list(request):
    tenders = Tender.objects.all()
    if request.GET.get('my') == '1' and request.user.is_authenticated:
        tenders = tenders.filter(user=request.user)
    tenders = tenders[:20]
    serializer = TenderListSerializer(tenders, many=True)
    return Response(serializer.data)


@api_view(['GET'])
def tender_detail(request, pk):
    try:
        tender = Tender.objects.get(pk=pk)
    except Tender.DoesNotExist:
        return Response({'error': 'Тендер не найден'}, status=status.HTTP_404_NOT_FOUND)

    serializer = TenderDetailSerializer(tender)
    return Response(serializer.data)


@api_view(['POST'])
@parser_classes([MultiPartParser])
def tender_upload(request):
    serializer = TenderUploadSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    tender = serializer.save()

    if request.user.is_authenticated:
        tender.user = request.user
        tender.save(update_fields=['user'])

    # Start AI pipeline in background thread
    thread = threading.Thread(target=process_tender, args=(tender.id,), daemon=True)
    thread.start()

    return Response(
        TenderDetailSerializer(tender).data,
        status=status.HTTP_201_CREATED,
    )


@api_view(['GET'])
def demo_pdf_list(request):
    seed_dir = Path(settings.BASE_DIR) / 'seed_data' / 'tenders'
    files = []
    if seed_dir.exists():
        for f in sorted(seed_dir.glob('*.pdf')):
            files.append({'name': f.stem.replace('_', ' ').title(), 'filename': f.name})
    return Response(files)


@api_view(['GET'])
def demo_pdf_download(request, filename):
    filepath = Path(settings.BASE_DIR) / 'seed_data' / 'tenders' / filename
    if not filepath.exists() or not filepath.suffix == '.pdf':
        return Response({'error': 'Файл не найден'}, status=status.HTTP_404_NOT_FOUND)
    return FileResponse(open(filepath, 'rb'), content_type='application/pdf', filename=filename)


@api_view(['POST'])
def monopoly_check(request):
    tender_ids = request.data.get('tender_ids', [])
    if len(tender_ids) < 2:
        return Response(
            {'error': 'Выберите минимум 2 тендера для проверки'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    tenders = Tender.objects.filter(id__in=tender_ids, status='completed')
    if tenders.count() < 2:
        return Response(
            {'error': 'Минимум 2 проанализированных тендера'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    tenders_data = [prepare_tender_data(t) for t in tenders]
    result = check_monopoly(tenders_data)
    return Response(result)


@api_view(['GET'])
def tender_export(request, pk):
    try:
        tender = Tender.objects.prefetch_related('requirements').get(pk=pk)
    except Tender.DoesNotExist:
        return Response({'error': 'Тендер не найден'}, status=status.HTTP_404_NOT_FOUND)

    if tender.status != 'completed':
        return Response({'error': 'Анализ ещё не завершён'}, status=status.HTTP_400_BAD_REQUEST)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    h = doc.add_heading('Анализ тендера — TendrAI', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(tender.title or 'Без названия', level=1)

    if tender.summary:
        s = tender.summary
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        for label, val in [
            ('Заказчик', s.get('customer', '')),
            ('Сумма', s.get('amount', '')),
            ('Дедлайн', s.get('deadline', '')),
            ('Категория', s.get('category', '')),
            ('Место поставки', s.get('delivery_location', '')),
        ]:
            if val:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(val)

    if tender.risk_score is not None:
        doc.add_heading(f'Оценка риска: {tender.risk_score}/100', level=2)
        if tender.risk_explanation:
            doc.add_paragraph(tender.risk_explanation)

    reqs = tender.requirements.all()
    if reqs:
        doc.add_heading(f'Требования ({reqs.count()})', level=2)
        cat_labels = {
            'qualification': 'Квалификация', 'technical': 'Технические',
            'financial': 'Финансовые', 'deadline': 'Сроки', 'document': 'Документы',
        }
        status_labels = {
            'critical': '[КРИТИЧНО]', 'warning': '[ВНИМАНИЕ]',
            'ok': '[OK]', 'needs_review': '[ПРОВЕРИТЬ]',
        }
        grouped = {}
        for r in reqs:
            grouped.setdefault(r.category, []).append(r)
        for cat, items in grouped.items():
            doc.add_heading(cat_labels.get(cat, cat), level=3)
            for r in items:
                prefix = status_labels.get(r.status, '')
                p = doc.add_paragraph(f'{prefix} {r.text}')
                if r.details:
                    doc.add_paragraph(r.details, style='List Bullet')

    if tender.pitfalls:
        doc.add_heading(f'Подводные камни ({len(tender.pitfalls)})', level=2)
        for pit in tender.pitfalls:
            p = doc.add_paragraph(pit.get('text', ''))
            p.runs[0].bold = True
            if pit.get('recommendation'):
                doc.add_paragraph(f'Рекомендация: {pit["recommendation"]}', style='List Bullet')

    if tender.technical_proposal:
        doc.add_heading('Техническое предложение (черновик)', level=2)
        for line in tender.technical_proposal.split('\n'):
            doc.add_paragraph(line)

    doc.add_paragraph('')
    footer = doc.add_paragraph('Сгенерировано TendrAI — AI-ассистент госзакупок РК')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = (tender.title or 'tender')[:50].replace(' ', '_')
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="TendrAI_{safe_title}.docx"'
    return response


@api_view(['POST'])
def analyze_goszakup(request):
    url = request.data.get('url', '').strip()
    if not url:
        return Response({'error': 'URL объявления не указан'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        pdf_bytes, filename = fetch_tender_pdf(url)
    except ValueError as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    except Exception:
        return Response(
            {'error': 'Не удалось скачать PDF с goszakup.gov.kz. Попробуйте загрузить файл вручную.'},
            status=status.HTTP_502_BAD_GATEWAY,
        )

    from django.core.files.base import ContentFile
    tender = Tender.objects.create()
    tender.pdf_file.save(filename, ContentFile(pdf_bytes), save=True)

    if request.user.is_authenticated:
        tender.user = request.user
        tender.save(update_fields=['user'])

    thread = threading.Thread(target=process_tender, args=(tender.id,), daemon=True)
    thread.start()

    return Response(TenderDetailSerializer(tender).data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
def search_goszakup(request):
    query = request.GET.get('q', '').strip()
    if not query:
        return Response({'results': [], 'total': 0})

    page = int(request.GET.get('page', 1))
    year = int(request.GET.get('year', 2026))
    result = goszakup_search(query, page=page, year=year)
    return Response(result)


def tender_stream(request, pk):
    """SSE endpoint for real-time progress updates."""
    def event_stream():
        prev_status = None
        prev_message = None
        while True:
            try:
                tender = Tender.objects.get(pk=pk)
            except Tender.DoesNotExist:
                yield f"data: {json.dumps({'error': 'not_found'})}\n\n"
                return

            current = {
                'status': tender.status,
                'message': tender.progress_message,
                'title': tender.title,
                'risk_score': tender.risk_score,
                'page_count': tender.page_count,
            }

            if tender.status != prev_status or tender.progress_message != prev_message:
                yield f"data: {json.dumps(current, ensure_ascii=False)}\n\n"
                prev_status = tender.status
                prev_message = tender.progress_message

            if tender.status in ('completed', 'failed'):
                return

            time.sleep(0.5)

    response = StreamingHttpResponse(
        event_stream(),
        content_type='text/event-stream',
    )
    response['Cache-Control'] = 'no-cache'
    response['X-Accel-Buffering'] = 'no'
    return response
